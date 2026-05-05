"""Word Document Handler — Create & edit .docx files"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime


class WordHandler:
    def __init__(self):
        self.output_dir = os.path.expanduser('~/Documents/PocketAI')
        os.makedirs(self.output_dir, exist_ok=True)

    def create(self, title='Untitled', content='', template='blank',
               sections=None, author='Pocket AI Office', **kwargs):
        """Create a new Word document"""
        doc = Document()

        # Set document properties
        doc.core_properties.author = author
        doc.core_properties.title = title

        # Apply template styling
        self._apply_template(doc, template)

        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_para = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph('')  # Spacer

        # Add sections or plain content
        if sections:
            for section in sections:
                heading = section.get('heading', '')
                body = section.get('content', '')
                level = section.get('level', 1)

                if heading:
                    doc.add_heading(heading, level=level)
                if body:
                    doc.add_paragraph(body)
        elif content:
            for paragraph in content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

        # Save with unique filename if busy
        base_filename = self._safe_filename(title)
        filename = base_filename + '.docx'
        filepath = os.path.join(self.output_dir, filename)
        
        counter = 1
        while os.path.exists(filepath):
            try:
                # Try saving - if it works, the file exists but isn't locked by Word
                # If it fails, it's either locked or we need a new name
                doc.save(filepath)
                break
            except PermissionError:
                filename = f"{base_filename}_{counter}.docx"
                filepath = os.path.join(self.output_dir, filename)
                counter += 1
            except Exception as e:
                # Fallback to a unique name if any other error
                filename = f"{base_filename}_{datetime.now().strftime('%H%M%S')}.docx"
                filepath = os.path.join(self.output_dir, filename)
                doc.save(filepath)
                break
        else:
            # File didn't exist yet
            doc.save(filepath)
        
        # Open the file automatically on the laptop
        try:
            os.startfile(filepath)
        except:
            pass

        return {
            'data': {'message': f'Document "{title}" created', 'path': filepath},
            'filesCreated': [{'name': filename, 'path': filepath, 'type': 'document',
                              'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                              'size': os.path.getsize(filepath)}],
        }

    def edit(self, path='', operations=None, **kwargs):
        """Edit an existing Word document"""
        if not os.path.exists(path):
            raise FileNotFoundError(f'Document not found: {path}')

        doc = Document(path)

        if operations:
            for op in operations:
                op_type = op.get('type')
                if op_type == 'add_paragraph':
                    doc.add_paragraph(op.get('text', ''))
                elif op_type == 'add_heading':
                    doc.add_heading(op.get('text', ''), level=op.get('level', 1))
                elif op_type == 'replace':
                    self._replace_text(doc, op.get('find', ''), op.get('replace', ''))
                elif op_type == 'add_image':
                    self._add_image(doc, op.get('image_path', ''), width=op.get('width', 6))
        try:
            doc.save(path)
        except PermissionError:
            # If original file is locked, save a copy
            base, ext = os.path.splitext(path)
            new_path = f"{base}_edited_{datetime.now().strftime('%H%M%S')}{ext}"
            doc.save(new_path)
            path = new_path
        except Exception as e:
            raise e
        
        # Open the file automatically
        try:
            os.startfile(path)
        except:
            pass

        return {'data': {'message': f'Document edited', 'path': path}}

    def add_image(self, path='', image_path='', width=6, **kwargs):
        """Specifically add an image to a document"""
        if not os.path.exists(path):
            raise FileNotFoundError(f'Document not found: {path}')
        
        doc = Document(path)
        self._add_image(doc, image_path, width)
        doc.save(path)
        
        return {'data': {'message': 'Image added to document', 'path': path}}

    def export_pdf(self, path='', **kwargs):
        """Export document info (actual PDF conversion needs LibreOffice)"""
        return {'data': {'message': 'PDF export queued', 'path': path,
                         'note': 'Requires LibreOffice for conversion'}}

    def _apply_template(self, doc, template):
        """Apply styling based on template name"""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        if template == 'formal':
            font.name = 'Times New Roman'
            font.size = Pt(12)
        elif template == 'modern':
            font.name = 'Arial'
            font.size = Pt(11)

    def _replace_text(self, doc, find, replace):
        for para in doc.paragraphs:
            if find in para.text:
                for run in para.runs:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)

    def _add_image(self, doc, image_path, width=6):
        if not image_path or not os.path.exists(image_path):
            return
        doc.add_picture(image_path, width=Inches(width))

    def _safe_filename(self, name):
        return "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).strip()
