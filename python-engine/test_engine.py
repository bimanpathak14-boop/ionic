import os
import sys

def test_everything():
    print("--- Pocket AI Engine Test ---")
    
    # 1. Check Libraries
    print("1. Checking Libraries...", end=" ")
    try:
        import PIL
        import pytesseract
        import docx
        print("✅ SUCCESS")
    except ImportError as e:
        print(f"❌ FAILED: {str(e)}")
        return

    # 2. Check Tesseract
    print("2. Checking Tesseract...", end=" ")
    tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    if os.path.exists(tesseract_path):
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        try:
            version = pytesseract.get_tesseract_version()
            print(f"✅ SUCCESS (Version: {version})")
        except Exception as e:
            print(f"❌ FAILED to run Tesseract: {str(e)}")
    else:
        print(f"❌ FAILED: Tesseract not found at {tesseract_path}")

    # 3. Create a Proof Word Document
    print("3. Creating Word Document...", end=" ")
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Pocket AI Office - System Test', 0)
        doc.add_paragraph('Hello! If you are reading this, it means your Python Automation Engine is WORKING PERFECTLY.')
        doc.add_paragraph(f'Tested on: {os.name} system')
        
        output_path = os.path.join(os.path.expanduser('~'), 'Desktop', 'Pocket_AI_Proof.docx')
        doc.save(output_path)
        print(f"✅ SUCCESS")
        print(f"\n✨ BINGO! Check your Desktop for 'Pocket_AI_Proof.docx'")
    except Exception as e:
        print(f"❌ FAILED to create Word file: {str(e)}")

if __name__ == "__main__":
    test_everything()
